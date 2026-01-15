from fpdf import FPDF
from docx import Document
import os
import re

class FormatService:
    
    # YARDIMCI: Türkçe karakterleri FPDF'in anlayacağı formata çevirir
    def fix_text(self, text):
        if not text: return ""
        text = str(text)
        
        # 1. HTML taglerini temizle (Regex ile)
        text = re.sub('<[^<]+?>', '', text)
        
        # 2. Türkçe Karakter Haritalama (Standart fontlar için)
        replacements = {
            'ğ': 'g', 'Ğ': 'G',
            'ü': 'u', 'Ü': 'U',
            'ş': 's', 'Ş': 'S',
            'ı': 'i', 'İ': 'I',
            'ö': 'o', 'Ö': 'O',
            'ç': 'c', 'Ç': 'C',
            '–': '-', '—': '-',
            '"': '"', '"': '"',
            "'": '\'', "'": '\''
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
            
        # 3. ASCII dışı karakterleri temizle (Çökmemesi için garanti)
        return text.encode('latin-1', 'replace').decode('latin-1')

    def generate_pdf(self, data, content_id, title):
        filename = f"summary_{content_id}.pdf"
        static_folder = os.path.join('app', 'static')
        
        # Klasör yoksa oluştur
        if not os.path.exists(static_folder):
            os.makedirs(static_folder)
            
        filepath = os.path.join(static_folder, filename)
        
        try:
            # PDF Oluşturucu Başlat
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # --- BAŞLIK ---
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(200, 10, txt=self.fix_text(title), ln=True, align='C')
            pdf.ln(10) # Boşluk
            
            # --- ÖZET ---
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(200, 10, txt="Summary", ln=True, align='L')
            
            pdf.set_font("Arial", size=11)
            # Multi_cell metni otomatik alt satıra indirir
            pdf.multi_cell(0, 10, txt=self.fix_text(data['summary']))
            pdf.ln(10)
            
            # --- QUIZ ---
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(200, 10, txt="Quiz Questions", ln=True, align='L')
            
            pdf.set_font("Arial", size=11)
            for i, q in enumerate(data['quiz']):
                question = f"{i+1}. {q['question']}"
                answer = f"Answer: {q['answer']}"
                
                pdf.multi_cell(0, 8, txt=self.fix_text(question))
                pdf.set_font("Arial", 'I', 11) # İtalik cevap
                pdf.multi_cell(0, 8, txt=self.fix_text(answer))
                pdf.set_font("Arial", size=11) # Normale dön
                pdf.ln(5)

            pdf.output(filepath)
            return filename
            
        except Exception as e:
            print(f"FPDF ERROR: {e}")
            raise e

    def generate_word(self, data, content_id, title):
        filename = f"summary_{content_id}.docx"
        static_folder = os.path.join('app', 'static')
        
        if not os.path.exists(static_folder):
            os.makedirs(static_folder)
            
        filepath = os.path.join(static_folder, filename)
        
        try:
            doc = Document()
            doc.add_heading(title, 0)
            
            doc.add_heading('Summary', level=1)
            doc.add_paragraph(data['summary'])
            
            doc.add_heading('Quiz Questions', level=1)
            for i, q in enumerate(data['quiz']):
                doc.add_paragraph(f"{i+1}. {q['question']}", style='List Number')
                doc.add_paragraph(f"Answer: {q['answer']}")
                
            doc.save(filepath)
            return filename
        except Exception as e:
            print(f"WORD ERROR: {e}")
            raise e